#!/usr/bin/env python3
"""Generate a Word document mirroring the Monty Tadier campaign website."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Color palette from the website CSS custom properties
BLUE_DEEP = RGBColor(0x0A, 0x22, 0x40)
BLUE_MID = RGBColor(0x1A, 0x4A, 0x7A)
BLUE_BRIGHT = RGBColor(0x1E, 0x6F, 0xBF)
AZURE = RGBColor(0x29, 0x79, 0xC2)
GOLD = RGBColor(0xC8, 0xA8, 0x4B)
CREAM = RGBColor(0xF4, 0xF0, 0xE8)
TEXT_COLOR = RGBColor(0x1A, 0x2A, 0x3A)
MUTED = RGBColor(0x5A, 0x70, 0x90)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SILVER_BG = "E8EEF5"

IMG_DIR = os.path.join(os.path.dirname(__file__), "img")

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = TEXT_COLOR
style.paragraph_format.space_after = Pt(4)

# -- Helpers --
def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(text, level=1, color=BLUE_DEEP, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    h = doc.add_heading(text, level=level)
    h.alignment = alignment
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Georgia'  # Closest to Cormorant Garamond
    return h

def add_para(text, bold=False, italic=False, color=TEXT_COLOR, size=Pt(11), alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = size
    run.font.name = 'Calibri'
    return p

def add_gold_rule():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— ◆ —")
    run.font.color.rgb = GOLD
    run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_section_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"— {text.upper()}")
    run.font.color.rgb = AZURE
    run.font.size = Pt(9)
    run.bold = True
    run.font.name = 'Calibri'

def set_row_height(row, height_cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{int(height_cm * 567)}" w:hRule="atLeast"/>')
    trPr.append(trHeight)

def remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def style_table_no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            remove_cell_borders(cell)


# ============================================================
# HEADER / TITLE BLOCK (dark navy background)
# ============================================================
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, "0A2240")
style_table_no_borders(table)

# Name
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("MONTY TADIER")
run.font.size = Pt(32)
run.font.color.rgb = CREAM
run.font.name = 'Georgia'
run.bold = True

# Gold accent
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run("Deputy for St Brelade")
run2.font.size = Pt(16)
run2.font.color.rgb = GOLD
run2.font.name = 'Georgia'
run2.italic = True

# Subtitle
p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(20)
run3 = p3.add_run("Jersey 2026 Election")
run3.font.size = Pt(12)
run3.font.color.rgb = CREAM
run3.font.name = 'Calibri'

# ============================================================
# BANNER IMAGE
# ============================================================
banner_path = os.path.join(IMG_DIR, "monfortbanner.png")
if os.path.exists(banner_path):
    doc.add_picture(banner_path, width=Inches(6.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    last_paragraph.paragraph_format.space_after = Pt(12)

# ============================================================
# HERO ANNOUNCEMENT (blue box)
# ============================================================
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, "0A2240")
style_table_no_borders(table)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("I'm standing for re-election as your Deputy for St Brelade")
run.font.size = Pt(18)
run.font.color.rgb = CREAM
run.font.name = 'Georgia'
run.bold = True

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(16)
run2 = p2.add_run("I'm ready to keep fighting for our community.")
run2.font.size = Pt(14)
run2.font.color.rgb = GOLD
run2.font.name = 'Georgia'
run2.bold = True

# Hero body text
hero_texts = [
    "For nearly two decades, I've had the privilege of representing this parish in the States Assembly. Since 2008, I've worked to be the kind of representative St Brelade deserves: one who shows up, speaks up, and never stops listening.",
    "Your trust over these 18 years has meant everything — and it has never made me complacent. If anything, the challenges facing Jersey today demand more determination, not less. I bring both the experience of what works and the energy to push for what must change.",
    "This election is about the next four years — and I want them shaped by you. In the coming weeks, I'll be knocking on doors across St Brelade to hear what matters most to you: what's working, what isn't, and what kind of Jersey you want to see. My full manifesto will be delivered to every home in the parish, but you can explore my vision and values right here.",
]
for t in hero_texts:
    add_para(t, color=TEXT_COLOR, size=Pt(11), space_after=Pt(10))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(12)
run = p.add_run("Your voice. Your parish. Your choice.")
run.font.size = Pt(15)
run.font.color.rgb = TEXT_COLOR
run.font.name = 'Georgia'
run.italic = True
run.bold = True

# ============================================================
# BANNER STRIP (blue bar with values)
# ============================================================
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, "1E6FBF")
style_table_no_borders(table)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)
values = ["Fairness", "Equality", "Affordable Healthcare", "Environmental Protection", "Arts & Culture", "Accountability"]
run = p.add_run("  ◆  ".join(values))
run.font.size = Pt(11)
run.font.color.rgb = WHITE
run.font.name = 'Georgia'
run.italic = True

# ============================================================
# MY VISION
# ============================================================
add_section_label("My Vision")
add_heading_styled("Positive Politics. Real Results.", level=1, color=BLUE_DEEP)

# Try to add the candidate photo
mypic_path = os.path.join(IMG_DIR, "mypic.jpg")
if os.path.exists(mypic_path):
    doc.add_picture(mypic_path, width=Inches(2.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)

add_para(
    "I believe Jersey works best when its government genuinely listens — to pensioners, young families, working people, and everyone in between. For over 16 years I have brought that commitment to the States Assembly, fighting for a fairer island even when it means challenging my own government.",
    color=MUTED, size=Pt(11), space_after=Pt(8)
)
add_para(
    "From securing sustainable arts funding to changing the law on medicinal cannabis, from protecting St Brelade's coastline from overdevelopment to pushing for affordable GP access — I measure success not in speeches, but in outcomes.",
    color=MUTED, size=Pt(11), space_after=Pt(12)
)

# Vision points
vision_points = [
    ("Longest-serving St Brelade Deputy", "First elected in 2008, re-elected four times — a record of consistent community service and trusted representation."),
    ("Reform Jersey & independent-minded", "A founding voice of Reform Jersey, committed to social-democratic values, green politics, and evidence-led decision-making."),
    ("Accessible to every islander", "Regular parish surgeries at St Aubin and Les Quennevais, always available by phone, email, and social media."),
]
for title, desc in vision_points:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("✓  ")
    run.font.color.rgb = AZURE
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = BLUE_DEEP
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.left_indent = Cm(1)
    run2 = p2.add_run(desc)
    run2.font.color.rgb = MUTED
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'

add_gold_rule()

# ============================================================
# KEY PRIORITIES (dark navy background section)
# ============================================================
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, "0A2240")
style_table_no_borders(table)

# Section header
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(2)
run = p.add_run("— KEY PRIORITIES")
run.font.color.rgb = GOLD
run.font.size = Pt(9)
run.bold = True
run.font.name = 'Calibri'

p2 = cell.add_paragraph()
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run("What I'm Fighting For")
run2.font.size = Pt(24)
run2.font.color.rgb = CREAM
run2.font.name = 'Georgia'
run2.bold = True

p3 = cell.add_paragraph()
p3.paragraph_format.space_after = Pt(16)
run3 = p3.add_run("Six areas where St Brelade — and Jersey — urgently needs principled, experienced leadership.")
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x9A, 0xAA, 0xBB)
run3.font.name = 'Calibri'

priorities = [
    ("01", "Affordable Healthcare", "GP access must be free for those who need it most. I've worked with oncology patients and medical experts to bring radiotherapy on-island, and changed the law to allow medicinal cannabis prescriptions — transforming lives that were otherwise living with unbearable pain."),
    ("02", "Housing & Cost of Living", "Young families are being priced out of the island they grew up in. We need stronger protections for renters, genuinely affordable housing, and targeted support for working households facing the highest cost of living in a generation."),
    ("03", "Environment & Coastline", "St Brelade's coastline — La Pulente, Portelet, St Aubin's Bay — is a shared treasure. I've led protests against coastal overdevelopment and proposed that Jersey's agricultural fields must be kept for farming, not ground-mounted solar or speculative development."),
    ("04", "Arts, Culture & Heritage", "Culture is not a luxury — it's economic and social infrastructure. I secured permanent, sustainable funding for the arts sector at 1% of government revenue expenditure, and successfully defended it against government attempts to water it down in 2024."),
    ("05", "Education & Languages", "As President of the Assemblée Parlementaire de la Francophonie (Jersey Branch), I've championed compulsory second-language study to GCSE level, reversing a 20-year policy that has left Jersey out of step with Europe and weakened our cultural identity."),
    ("06", "Accountability & Scrutiny", "Currently Chair of the Economic and International Affairs Scrutiny Panel, I hold government to account on the decisions that shape islanders' daily lives — from public finances to international trade. Transparency is not optional."),
]

for num, title, body in priorities:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(num)
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x3A, 0x50, 0x60)
    run.font.name = 'Georgia'
    run.bold = True
    run = p.add_run(f"  {title}")
    run.font.size = Pt(14)
    run.font.color.rgb = CREAM
    run.font.name = 'Georgia'
    run.bold = True

    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    p2.paragraph_format.left_indent = Cm(0.5)
    run2 = p2.add_run(body)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x9A, 0xAA, 0xBB)
    run2.font.name = 'Calibri'

# Add bottom padding
p_end = cell.add_paragraph()
p_end.paragraph_format.space_after = Pt(16)

# ============================================================
# THE CHALLENGES WE FACE (silver background)
# ============================================================
add_section_label("The Challenges We Face")
add_heading_styled("Jersey at a Crossroads", level=1, color=BLUE_DEEP)

add_para(
    "The cost of living, the housing crisis, pressure on our health service, and the need to protect our environment — these aren't abstract political issues. They are the lived reality of St Brelade families every single day.",
    color=MUTED, size=Pt(11), space_after=Pt(10)
)

# Quote
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
# Gold left border via shading + left border
tc = cell._tc
tcPr = tc.get_or_add_tcPr()
tcBorders = parse_xml(
    f'<w:tcBorders {nsdecls("w")}>'
    '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:left w:val="single" w:sz="18" w:space="0" w:color="C8A84B"/>'
    '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tcBorders>'
)
tcPr.append(tcBorders)
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
run = p.add_run('"Young families must be given hope of staying in Jersey and having a good quality of life. We must tackle the cost of living urgently."')
run.font.size = Pt(13)
run.font.color.rgb = BLUE_DEEP
run.font.name = 'Georgia'
run.italic = True

# Issues
issues = [
    ("Housing Crisis", "Rents and house prices have risen far beyond ordinary wages. Tenants need stronger protections and more affordable supply — now."),
    ("NHS-style Healthcare Access", "No islander should avoid seeing a GP because they cannot afford it. Waiting times and cost barriers must be tackled together."),
    ("Coastal Overdevelopment", "Our Coastal National Park and shoreline must be protected from speculative development. The coastline belongs to all islanders — permanently."),
    ("Government Accountability", "Too many decisions are made behind closed doors. Islanders deserve transparent, evidence-led government that answers for its choices."),
]

for i, (title, body) in enumerate(issues):
    accent = GOLD if i % 2 == 1 else AZURE
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("■ ")
    run.font.color.rgb = accent
    run.font.size = Pt(11)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = BLUE_DEEP
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.left_indent = Cm(0.75)
    run2 = p2.add_run(body)
    run2.font.color.rgb = MUTED
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'

# Stat cards as a table
doc.add_paragraph()  # spacer
stats_table = doc.add_table(rows=2, cols=2)
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER

stats = [
    ("16+", "years serving St Brelade", AZURE),
    ("5×", "elected to the States Assembly", GOLD),
    ("£11.6m", "secured annually for Jersey arts & culture", GOLD),
    ("800+", "islanders backed Monty's petition to protect agricultural land", AZURE),
]

for idx, (num, label, accent) in enumerate(stats):
    row = idx // 2
    col = idx % 2
    cell = stats_table.cell(row, col)
    set_cell_shading(cell, "FFFFFF")

    # Top border accent
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    border_hex = "2979C2" if accent == AZURE else "C8A84B"
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="18" w:space="0" w:color="{border_hex}"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(num)
    run.font.size = Pt(24)
    run.font.color.rgb = BLUE_DEEP
    run.font.name = 'Georgia'
    run.bold = True

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(label)
    run2.font.size = Pt(9)
    run2.font.color.rgb = MUTED
    run2.font.name = 'Calibri'

# Recent Achievements
add_gold_rule()
add_heading_styled("Recent Achievements", level=2, color=BLUE_DEEP)

achievements = [
    "Changed the law to allow GPs to prescribe medicinal cannabis",
    "Secured permanent arts & culture funding at 1% of revenue expenditure",
    "Led coastal protection protests at La Pulente against overdevelopment",
    "Advocated for on-island radiotherapy provision with oncology patients",
    "Championed compulsory language study to GCSE to reverse 20-year decline",
]
for a in achievements:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("◆  ")
    run.font.color.rgb = GOLD
    run.font.size = Pt(9)
    run = p.add_run(a)
    run.font.color.rgb = TEXT_COLOR
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'

add_gold_rule()

# ============================================================
# ABOUT MONTY
# ============================================================
add_section_label("About Monty")
add_heading_styled("Jersey-Born. St Brelade Through & Through.", level=1, color=BLUE_DEEP)

# St Aubin image
staubin_path = os.path.join(IMG_DIR, "StAubin.jpg")
if os.path.exists(staubin_path):
    doc.add_picture(staubin_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

about_texts = [
    "Montfort (Monty) Tadier was born in St Helier on 28 November 1979, grew up in St Brelade, and attended La Moye Primary School, Les Quennevais and Hautlieu Secondary Schools before reading Modern Languages at the University of Sheffield.",
    "He identifies with green, social-democratic and socialist politics. A co-founder of Reform Jersey — registered in 2014 — he has been a consistent, principled voice for fairness and equality throughout his public life. He is also a qualified Jersey tour guide, musician and community organiser.",
    "Monty holds regular constituency surgeries at The Boathouse in St Aubin and the Communicare Café in Les Quennevais, keeping him directly connected to the people he represents.",
]
for t in about_texts:
    add_para(t, color=MUTED, size=Pt(11), space_after=Pt(10))

# Facts
facts = [
    ("Education:", "BA (Hons) Modern Languages, University of Sheffield"),
    ("Party:", "Reform Jersey (co-founder)"),
    ("Current role:", "Chair, Economic & International Affairs Scrutiny Panel"),
    ("Also:", "Qualified Jersey tour guide · Member, Assemblée Parlementaire de la Francophonie"),
    ("Contact:", "montfort.tadier@gov.je · 07797 844 358"),
]
for label, value in facts:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run("◆  ")
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)
    run = p.add_run(label + " ")
    run.bold = True
    run.font.color.rgb = BLUE_DEEP
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run = p.add_run(value)
    run.font.color.rgb = MUTED
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

# Record of Service
doc.add_paragraph()
add_heading_styled("Record of Service", level=2, color=BLUE_DEEP)

roles = [
    ("2008 – present", "Deputy for St Brelade (5 consecutive terms)"),
    ("2024 – present", "Chair, Economic & International Affairs Scrutiny Panel"),
    ("2019 – 2022", "Assistant Minister for Economic Development, Tourism, Sport & Culture"),
    ("2020", "Assistant Minister for Environment"),
    ("2008 – 2013", "Education & Home Affairs Scrutiny Panel"),
    ("2015 – present", "President, Assemblée Parlementaire de la Francophonie (Jersey)"),
]

roles_table = doc.add_table(rows=len(roles), cols=2)
roles_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (year, role) in enumerate(roles):
    cell_y = roles_table.cell(i, 0)
    cell_r = roles_table.cell(i, 1)
    remove_cell_borders(cell_y)
    remove_cell_borders(cell_r)

    p = cell_y.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(year)
    run.font.size = Pt(9)
    run.font.color.rgb = GOLD
    run.bold = True
    run.font.name = 'Calibri'

    p2 = cell_r.paragraphs[0]
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(role)
    run2.font.size = Pt(10)
    run2.font.color.rgb = MUTED
    run2.font.name = 'Calibri'

# Quote box
doc.add_paragraph()
table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, "0A2240")
tc = cell._tc
tcPr = tc.get_or_add_tcPr()
tcBorders = parse_xml(
    f'<w:tcBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="4" w:space="0" w:color="C8A84B"/>'
    '  <w:left w:val="single" w:sz="4" w:space="0" w:color="C8A84B"/>'
    '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="C8A84B"/>'
    '  <w:right w:val="single" w:sz="4" w:space="0" w:color="C8A84B"/>'
    '</w:tcBorders>'
)
tcPr.append(tcBorders)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(6)
run = p.add_run('"I have always sought to be an active and accessible constituency representative — a strong voice for pensioners, young people and working families in St Brelade."')
run.font.size = Pt(12)
run.font.color.rgb = CREAM
run.font.name = 'Georgia'
run.italic = True

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(14)
run2 = p2.add_run("— Monty Tadier, Deputy for St Brelade")
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(0x6A, 0x7A, 0x8A)
run2.font.name = 'Calibri'

# ============================================================
# FOOTER
# ============================================================
add_gold_rule()

table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.cell(0, 0)
set_cell_shading(cell, "0A2240")
style_table_no_borders(table)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Deputy Montfort ")
run.font.size = Pt(18)
run.font.color.rgb = CREAM
run.font.name = 'Georgia'
run.bold = True
run = p.add_run("TADIER")
run.font.size = Pt(18)
run.font.color.rgb = GOLD
run.font.name = 'Georgia'
run.bold = True

footer_lines = [
    "Deputy for St Brelade",
    "Chair of the Economic and International Affairs Scrutiny Committee",
    "Président de l'Assemblée Parlementaire de la Francophonie",
]
for line in footer_lines:
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run(line)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x9A, 0xAA, 0xBB)
    run2.font.name = 'Calibri'

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(12)
p3.paragraph_format.space_after = Pt(4)
run3 = p3.add_run("© 2026 · Promoted by Monty Tadier · St Brelade, Jersey, Channel Islands")
run3.font.size = Pt(8)
run3.font.color.rgb = RGBColor(0x5A, 0x6A, 0x7A)
run3.font.name = 'Calibri'

p4 = cell.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_after = Pt(16)
run4 = p4.add_run("Designed and created by Coderra.je")
run4.font.size = Pt(8)
run4.font.color.rgb = GOLD
run4.font.name = 'Calibri'

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), "Monty_Tadier_Campaign.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
